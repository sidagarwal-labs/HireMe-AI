from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_resume_doc() -> Document:
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Default font for Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    return doc


def add_name_header(doc: Document, name: str, contact_line: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(16)

    p2 = doc.add_paragraph(contact_line)
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # small spacer
    doc.add_paragraph("")


def add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(12)


def add_paragraph(doc: Document, text: str, *, align_center: bool = False) -> None:
    p = doc.add_paragraph(text)
    if align_center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for b in bullets:
        b = (b or "").strip()
        if b:
            doc.add_paragraph(b, style="List Bullet")

def add_job_entry(
    doc: Document,
    job_title: str,
    company: str,
    start_date: str,
    end_date: str,
) -> None:
    """
    Formats:
    Job Title - Company
    Start Date - End Date
    """
    # Line 1: Title + Company (bold)
    p = doc.add_paragraph()
    r = p.add_run(f"{job_title} - {company}".strip(" -"))
    r.bold = True
    p.paragraph_format.space_after = Pt(0)

    # Line 2: Dates (italic)
    dates = f"{start_date} - {end_date}".strip(" -")
    p2 = doc.add_paragraph()
    r2 = p2.add_run(dates)
    r2.italic = True
    p2.paragraph_format.space_after = Pt(2)


def add_education_entry(
    doc: Document,
    degree: str,
    school: str,
    start_date: str,
    end_date: str,
    details: list[str] | None = None,
) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{degree} - {school}".strip(" -"))
    r.bold = True
    p.paragraph_format.space_after = Pt(0)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"{start_date} - {end_date}".strip(" -"))
    r2.italic = True
    p2.paragraph_format.space_after = Pt(2)

    if details:
        add_bullets(doc, details)           

def save_doc(doc: Document, out_path: str) -> None:
    doc.save(out_path)
